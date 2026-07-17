import streamlit as st
import pandas as pd
import plotly.express as px
from docx import Document
from io import BytesIO
from datetime import datetime
import base64


def get_base64_image(image_path):
    with open(image_path, "rb") as image_file:
        return base64.b64encode(
            image_file.read()
        ).decode()
# ==========================================================
# LOAD EXTERNAL CSS
# ==========================================================

def load_css(file_path):
    with open(file_path, "r", encoding="utf-8") as css_file:
        st.markdown(
            f"<style>{css_file.read()}</style>",
            unsafe_allow_html=True
        )
        
# ==========================================================
# PAGE CONFIGURATION
# ==========================================================

st.set_page_config(
    page_title="Pharma Compliance Dashboard",
    page_icon="💊",
    layout="wide"
)


# Load external CSS
load_css("styles/style.css")
# ==========================================================
# PROFESSIONAL DASHBOARD HEADER
# ==========================================================

logo_base64 = get_base64_image(
    "images/image.png"
)

header_html = f"""
<div class="pharma-header">
<div class="pharma-title">
<img
src="data:image/png;base64,{logo_base64}"
style="width:60px; height:60px; max-width:60px; max-height:60px; object-fit:cover; border-radius:50%;"
>
<span>Pharma Compliance Analytics Dashboard</span>
</div>
<div class="pharma-subtitle">
Compliance Monitoring &nbsp; | &nbsp; Risk Analytics &nbsp; | &nbsp; Management Decision Intelligence
</div>
</div>
"""

st.markdown(
    header_html,
    unsafe_allow_html=True
)
# ==========================================================
# LOAD DATASET
# ==========================================================

@st.cache_data
def load_data():
    data = pd.read_csv("data/pharma_compliance_cleaned.csv")

    # Convert date columns
    data["Audit_Date"] = pd.to_datetime(
        data["Audit_Date"],
        errors="coerce"
    )

    data["Closure_Date"] = pd.to_datetime(
        data["Closure_Date"],
        errors="coerce"
    )

    return data


df = load_data()


# ==========================================================
# SIDEBAR FILTERS
# ==========================================================

st.sidebar.header("🔍 Dashboard Filters")


company = st.sidebar.selectbox(
    "Select Company",
    ["All"] + sorted(df["Company"].dropna().unique().tolist())
)


department = st.sidebar.selectbox(
    "Select Department",
    ["All"] + sorted(df["Department"].dropna().unique().tolist())
)


plant = st.sidebar.selectbox(
    "Select Plant",
    ["All"] + sorted(df["Plant"].dropna().unique().tolist())
)


severity = st.sidebar.selectbox(
    "Select Severity",
    ["All"] + sorted(df["Severity"].dropna().unique().tolist())
)


# ==========================================================
# APPLY FILTERS
# ==========================================================

filtered_df = df.copy()


if company != "All":
    filtered_df = filtered_df[
        filtered_df["Company"] == company
    ]


if department != "All":
    filtered_df = filtered_df[
        filtered_df["Department"] == department
    ]


if plant != "All":
    filtered_df = filtered_df[
        filtered_df["Plant"] == plant
    ]


if severity != "All":
    filtered_df = filtered_df[
        filtered_df["Severity"] == severity
    ]


# ==========================================================
# KPI CALCULATIONS
# ==========================================================

total_cases = len(filtered_df)


average_risk = (
    round(filtered_df["Risk_Score"].mean(), 2)
    if total_cases > 0
    else 0
)


critical_cases = (
    filtered_df["Severity"] == "Critical"
).sum()


open_capa = (
    filtered_df["CAPA_Status"] == "Open"
).sum()


closed_capa = (
    filtered_df["CAPA_Status"] == "Closed"
).sum()


# ==========================================================
# KPI CARDS
# ==========================================================

st.subheader("📌 Compliance Overview")

col1, col2, col3, col4, col5 = st.columns(5)


def kpi_card(title, value):
    card_html = (
        '<div class="kpi-card">'
        f'<div class="kpi-title">{title}</div>'
        f'<div class="kpi-value">{value}</div>'
        '</div>'
    )

    st.markdown(
        card_html,
        unsafe_allow_html=True
    )


with col1:
    kpi_card(
        "📋 Total Cases",
        f"{total_cases:,}"
    )

with col2:
    kpi_card(
        "⚠️ Average Risk",
        f"{average_risk:.2f}"
    )

with col3:
    kpi_card(
        "🚨 Critical Cases",
        f"{critical_cases:,}"
    )

with col4:
    kpi_card(
        "📂 Open CAPA",
        f"{open_capa:,}"
    )

with col5:
    kpi_card(
        "✅ Closed CAPA",
        f"{closed_capa:,}"
    )
# ==========================================================
# MANAGEMENT DECISION INSIGHT CALCULATIONS
# ==========================================================

department_name_map = {
    "QA": "Quality Assurance",
    "QC": "Quality Control",
    "MFG": "Manufacturing",
    "PROD": "Production",
    "ENG": "Engineering",
    "WH": "Warehouse"
}


# --------------------------
# Highest Risk Department
# --------------------------

management_department_risk = (
    filtered_df
    .groupby("Department")["Risk_Score"]
    .mean()
    .sort_values(ascending=False)
)

if not management_department_risk.empty:

    highest_risk_department = (
        management_department_risk.index[0]
    )

    highest_risk_score = round(
        management_department_risk.iloc[0],
        2
    )

    highest_risk_department_display = (
        department_name_map.get(
            highest_risk_department,
            highest_risk_department
        )
    )

else:
    highest_risk_department_display = "N/A"
    highest_risk_score = 0


# --------------------------
# Top Root Cause
# --------------------------

root_cause_counts = (
    filtered_df["Root_Cause"]
    .dropna()
    .value_counts()
)

top_root_cause = (
    root_cause_counts.index[0]
    if not root_cause_counts.empty
    else "N/A"
)


# --------------------------
# Most Common Deviation
# --------------------------

deviation_type_counts = (
    filtered_df["Deviation_Type"]
    .dropna()
    .value_counts()
)

top_deviation = (
    deviation_type_counts.index[0]
    if not deviation_type_counts.empty
    else "N/A"
)


# --------------------------
# Slowest Closure Department
# --------------------------

management_capa = filtered_df.copy()

management_capa["Closure_Days"] = (
    management_capa["Closure_Date"]
    - management_capa["Audit_Date"]
).dt.days

management_capa = management_capa[
    management_capa["Closure_Days"].notna()
    & (management_capa["Closure_Days"] >= 0)
]

management_closure = (
    management_capa
    .groupby("Department")["Closure_Days"]
    .mean()
    .sort_values(ascending=False)
)

if not management_closure.empty:

    slowest_department = (
        management_closure.index[0]
    )

    slowest_closure_days = round(
        management_closure.iloc[0],
        1
    )

    slowest_department_display = (
        department_name_map.get(
            slowest_department,
            slowest_department
        )
    )
else:
    slowest_department_display = "N/A"
    slowest_closure_days = 0

st.markdown(
    "<div style='height: 10px;'></div>",
    unsafe_allow_html=True
)

st.subheader("🎯 Management Decision Insights")
# ==========================================================
# DISPLAY MANAGEMENT INSIGHT CARDS
# ==========================================================

def insight_card(title, value, detail=""):
    card_html = (
        '<div style="border:1px solid #E5E7EB;'
        'border-radius:12px;'
        'padding:20px;'
        'height:180px;'
        'box-sizing:border-box;'
        'background-color:#FFFFFF;'
        'box-shadow:0 2px 6px rgba(0,0,0,0.08);">'

        f'<div style="font-size:15px;'
        'color:#555555;'
        'margin-bottom:12px;">'
        f'{title}</div>'

        f'<div style="font-size:24px;'
        'font-weight:600;'
        'color:#1F2937;'
        'line-height:1.3;'
        'word-wrap:break-word;">'
        f'{value}</div>'

        f'<div style="font-size:14px;'
        'color:#16803A;'
        'margin-top:12px;">'
        f'{detail}</div>'

        '</div>'
    )

    st.markdown(
        card_html,
        unsafe_allow_html=True
    )


col1, col2, col3, col4 = st.columns(4)


with col1:
    insight_card(
        "🚨 Highest Risk Department",
        highest_risk_department_display,
        f"Risk Score: {highest_risk_score}"
    )


with col2:
    insight_card(
        "🔍 Top Root Cause",
        top_root_cause
    )


with col3:
    insight_card(
        "⚠️ Most Common Deviation",
        top_deviation
    )


with col4:
    insight_card(
        "⏱️ Slowest Closure Department",
        slowest_department_display,
        f"Average Closure: {slowest_closure_days} Days"
    )


# ==========================================================
# DEPARTMENT & SEVERITY ANALYSIS
# ==========================================================

st.divider()

col1, col2 = st.columns(2, gap="medium")


# --------------------------
# Department Analysis
# --------------------------

with col1:

    with st.container(border=True):

        department_counts = (
            filtered_df["Department"]
            .value_counts()
            .reset_index()
        )

        department_counts.columns = [
            "Department",
            "Cases"
        ]

        fig_department = px.bar(
            department_counts,
            x="Department",
            y="Cases",
            text="Cases",
            title="Department-wise Compliance Cases",
            color_discrete_sequence=["#4F81BD"],
            template="plotly_white"
        )

        fig_department.update_traces(
            textposition="outside"
        )

        fig_department.update_layout(
            height=430,
            title=dict(
                text="Department-wise Compliance Cases",
                x=0.5,
                xanchor="center",
                font=dict(size=20)
            ),
            xaxis_title="Department",
            yaxis_title="Number of Cases",
            margin=dict(
                l=40,
                r=30,
                t=70,
                b=50
            )
        )

        st.plotly_chart(
    fig_department,
    width="stretch",
    config={
        "responsive": True,
        "displayModeBar": False
    }
)


# --------------------------
# Severity Analysis
# --------------------------

with col2:

    with st.container(border=True):

        severity_counts = (
            filtered_df["Severity"]
            .value_counts()
            .reset_index()
        )

        severity_counts.columns = [
            "Severity",
            "Cases"
        ]

        fig_severity = px.pie(
            severity_counts,
            names="Severity",
            values="Cases",
            hole=0.45,
            title="Severity Distribution",
            template="plotly_white"
        )

        fig_severity.update_layout(
            height=430,
            title=dict(
                text="Severity Distribution",
                x=0.5,
                xanchor="center",
                font=dict(size=20)
            ),
            margin=dict(
                l=40,
                r=30,
                t=70,
                b=50
            ),
            legend=dict(
                orientation="h",
                yanchor="bottom",
                y=-0.15,
                xanchor="center",
                x=0.5
            )
        )

        st.plotly_chart(
    fig_severity,
    width="stretch",
    config={
        "responsive": True,
        "displayModeBar": False
    }
)

# ==========================================================
# CAPA & ROOT CAUSE ANALYSIS
# ==========================================================

st.divider()

col1, col2 = st.columns(2)


# --------------------------
# CAPA Status
# --------------------------


with col1:

    with st.container(border=True):

        capa_counts = (
            filtered_df["CAPA_Status"]
            .value_counts()
            .reset_index()
        )

        capa_counts.columns = [
            "CAPA Status",
            "Cases"
        ]

        fig_capa = px.pie(
            capa_counts,
            names="CAPA Status",
            values="Cases",
            hole=0.45,
            title="CAPA Status",
            template="plotly_white"
        )

        fig_capa.update_layout(
            height=430,

            title=dict(
                text="CAPA Status",
                x=0.5,
                xanchor="center",
                font=dict(size=20)
            ),

            margin=dict(
                l=40,
                r=30,
                t=70,
                b=50
            ),

            legend=dict(
                orientation="h",
                yanchor="bottom",
                y=-0.15,
                xanchor="center",
                x=0.5
            )
        )

        st.plotly_chart(
            fig_capa,
            width="stretch",
            config={
                "responsive": True,
                "displayModeBar": False
            }
        )
# --------------------------
# Root Cause Analysis
# --------------------------

with col2:

    with st.container(border=True):

        root_counts = (
            filtered_df["Root_Cause"]
            .value_counts()
            .reset_index()
        )

        root_counts.columns = [
            "Root Cause",
            "Cases"
        ]

        fig_root = px.bar(
            root_counts,
            x="Root Cause",
            y="Cases",
            text="Cases",
            title="Root Cause Analysis",
            color_discrete_sequence=["#2E8B57"],
            template="plotly_white"
        )

        fig_root.update_traces(
            textposition="outside"
        )

        fig_root.update_layout(
            height=430,

            title=dict(
                text="Root Cause Analysis",
                x=0.5,
                xanchor="center",
                font=dict(size=20)
            ),

            xaxis_title="Root Cause",
            yaxis_title="Number of Cases",

            margin=dict(
                l=40,
                r=30,
                t=70,
                b=50
            )
        )

        st.plotly_chart(
            fig_root,
            width="stretch",
            config={
                "responsive": True,
                "displayModeBar": False
            }
        )
# ==========================================================
# PLANT & PRODUCT ANALYSIS
# ==========================================================

st.divider()

col1, col2 = st.columns(2)


# --------------------------
# Plant Analysis
# --------------------------

with col1:

    with st.container(border=True):

        plant_counts = (
            filtered_df["Plant"]
            .value_counts()
            .reset_index()
        )

        plant_counts.columns = [
            "Plant",
            "Cases"
        ]

        fig_plant = px.bar(
            plant_counts,
            x="Plant",
            y="Cases",
            text="Cases",
            title="Plant-wise Compliance Cases",
            color_discrete_sequence=["#F39C12"],
            template="plotly_white"
        )

        fig_plant.update_traces(
            textposition="outside"
        )

        fig_plant.update_layout(
            height=430,

            title=dict(
                text="Plant-wise Compliance Cases",
                x=0.5,
                xanchor="center",
                font=dict(size=20)
            ),

            xaxis_title="Plant",
            yaxis_title="Number of Cases",

            margin=dict(
                l=40,
                r=30,
                t=70,
                b=50
            )
        )

        st.plotly_chart(
            fig_plant,
            width="stretch",
            config={
                "responsive": True,
                "displayModeBar": False
            }
        )

# --------------------------
# Product Analysis
# --------------------------

with col2:

    with st.container(border=True):

        product_counts = (
            filtered_df["Product"]
            .value_counts()
            .reset_index()
        )

        product_counts.columns = [
            "Product",
            "Cases"
        ]

        fig_product = px.bar(
            product_counts,
            x="Product",
            y="Cases",
            text="Cases",
            title="Product-wise Compliance Cases",
            color_discrete_sequence=["#8E44AD"],
            template="plotly_white"
        )

        fig_product.update_traces(
            textposition="outside"
        )

        fig_product.update_layout(
            height=430,

            title=dict(
                text="Product-wise Compliance Cases",
                x=0.5,
                xanchor="center",
                font=dict(size=20)
            ),

            xaxis_title="Product",
            yaxis_title="Number of Cases",

            margin=dict(
                l=40,
                r=30,
                t=70,
                b=50
            )
        )

        st.plotly_chart(
            fig_product,
            width="stretch",
            config={
                "responsive": True,
                "displayModeBar": False
            }
        )
# ==========================================================
# DEVIATION TYPE ANALYSIS
# ==========================================================

st.divider()

with st.container(border=True):

    deviation_counts = (
        filtered_df["Deviation_Type"]
        .value_counts()
        .reset_index()
    )

    deviation_counts.columns = [
        "Deviation Type",
        "Cases"
    ]

    fig_deviation = px.bar(
        deviation_counts,
        x="Deviation Type",
        y="Cases",
        text="Cases",
        title="Deviation Type Analysis",
        color_discrete_sequence=["#E67E22"],
        template="plotly_white"
    )

    fig_deviation.update_traces(
        textposition="outside"
    )

    fig_deviation.update_layout(
        height=430,

        title=dict(
            text="Deviation Type Analysis",
            x=0.5,
            xanchor="center",
            font=dict(size=20)
        ),

        xaxis_title="Deviation Type",
        yaxis_title="Number of Cases",

        margin=dict(
            l=50,
            r=30,
            t=70,
            b=60
        )
    )

    st.plotly_chart(
        fig_deviation,
        width="stretch",
        config={
            "responsive": True,
            "displayModeBar": False
        }
    )

# ==========================================================
# RISK ANALYSIS BY DEPARTMENT
# ==========================================================

st.divider()

with st.container(border=True):

    department_risk = (
        filtered_df
        .groupby(
            "Department",
            as_index=False
        )["Risk_Score"]
        .mean()
    )

    department_risk["Risk_Score"] = (
        department_risk["Risk_Score"]
        .round(2)
    )

    department_risk = (
        department_risk
        .sort_values(
            by="Risk_Score",
            ascending=False
        )
    )

    fig_risk = px.bar(
        department_risk,
        x="Department",
        y="Risk_Score",
        text="Risk_Score",
        color="Risk_Score",
        color_continuous_scale="Reds",
        title="Department-wise Risk Analysis",
        template="plotly_white"
    )

    fig_risk.update_traces(
        textposition="outside"
    )

    fig_risk.update_layout(
        height=430,

        title=dict(
            text="Department-wise Risk Analysis",
            x=0.5,
            xanchor="center",
            font=dict(size=20)
        ),

        xaxis_title="Department",
        yaxis_title="Average Risk Score",

        coloraxis_showscale=False,

        margin=dict(
            l=50,
            r=30,
            t=70,
            b=60
        )
    )

    st.plotly_chart(
        fig_risk,
        width="stretch",
        config={
            "responsive": True,
            "displayModeBar": False
        }
    )
# ==========================================================
# CAPA PERFORMANCE ANALYSIS
# ==========================================================

st.divider()
st.subheader("⏱️ CAPA Performance Analysis")

# Create a copy for closure analysis
capa_analysis = filtered_df.copy()

# Calculate case closure time
capa_analysis["Closure_Days"] = (
    capa_analysis["Closure_Date"]
    - capa_analysis["Audit_Date"]
).dt.days

# Keep records with valid, non-negative closure time
closed_cases_df = capa_analysis[
    capa_analysis["Closure_Days"].notna()
    & (capa_analysis["Closure_Days"] >= 0)
].copy()

# Calculate average closure time
average_closure_days = (
    round(closed_cases_df["Closure_Days"].mean(), 1)
    if not closed_cases_df.empty
    else 0
)

# Display closure KPI
# Display Average Closure Time KPI

# Display Average Closure Time KPI in a compact column

kpi_col, empty_col = st.columns([1, 3])

with kpi_col:
    closure_kpi_html = (
        '<div class="kpi-card">'
        '<div class="kpi-title">⏱️ Average Case Closure Time</div>'
        f'<div class="kpi-value">{average_closure_days} Days</div>'
        '</div>'
    )

    st.markdown(
        closure_kpi_html,
        unsafe_allow_html=True
    )

# ==========================================================
# DEPARTMENT-WISE AVERAGE CLOSURE TIME
# ==========================================================

with st.container(border=True):

    department_closure = (
        closed_cases_df
        .groupby(
            "Department",
            as_index=False
        )["Closure_Days"]
        .mean()
    )

    department_closure["Closure_Days"] = (
        department_closure["Closure_Days"]
        .round(1)
    )

    department_closure = (
        department_closure
        .sort_values(
            by="Closure_Days",
            ascending=False
        )
    )

    fig_closure = px.bar(
        department_closure,
        x="Department",
        y="Closure_Days",
        text="Closure_Days",
        title="Department-wise Average Closure Time",
        color_discrete_sequence=["#E67E22"],
        template="plotly_white"
    )

    fig_closure.update_traces(
        textposition="outside"
    )

    fig_closure.update_layout(
        height=430,

        title=dict(
            text="Department-wise Average Closure Time",
            x=0.5,
            xanchor="center",
            font=dict(size=20)
        ),

        xaxis_title="Department",
        yaxis_title="Average Closure Days",

        margin=dict(
            l=50,
            r=30,
            t=70,
            b=60
        )
    )

    st.plotly_chart(
        fig_closure,
        width="stretch",
        config={
            "responsive": True,
            "displayModeBar": False
        }
    )
  
# ==========================================================
# MONTHLY COMPLIANCE TREND
# ==========================================================

st.divider()

st.subheader("📈 Monthly Compliance Trend")

with st.container(border=True):

    # Create a copy of filtered data
    monthly_df = filtered_df.copy()

    # Convert Audit Date to datetime
    monthly_df["Audit_Date"] = pd.to_datetime(
        monthly_df["Audit_Date"],
        errors="coerce"
    )

    # Remove records with invalid dates
    monthly_df = monthly_df.dropna(
        subset=["Audit_Date"]
    )

    # Calculate monthly compliance cases
    monthly_cases = (
        monthly_df
        .groupby(
            monthly_df["Audit_Date"].dt.to_period("M")
        )
        .size()
        .reset_index(name="Cases")
    )

    # Convert Period to string for Plotly
    monthly_cases["Audit_Date"] = (
        monthly_cases["Audit_Date"]
        .astype(str)
    )

    # Create line chart
    fig_monthly = px.line(
        monthly_cases,
        x="Audit_Date",
        y="Cases",
        markers=True,
        title="Monthly Compliance Trend",
        template="plotly_white"
    )

    # Style line and markers
    fig_monthly.update_traces(
        line=dict(
            width=3
        ),
        marker=dict(
            size=8
        )
    )

    # Update chart layout
    fig_monthly.update_layout(
        height=430,

        title=dict(
            text="Monthly Compliance Trend",
            x=0.5,
            xanchor="center",
            font=dict(size=20)
        ),

        xaxis_title="Month",
        yaxis_title="Number of Compliance Cases",

        margin=dict(
            l=50,
            r=30,
            t=70,
            b=60
        )
    )

    # Display responsive chart
    st.plotly_chart(
        fig_monthly,
        width="stretch",
        config={
            "responsive": True,
            "displayModeBar": False
        }
    )


# ==========================================================
# WORD REPORT GENERATION
# ==========================================================

st.divider()

st.subheader("📄 Compliance Analytics Report")


document = Document()


document.add_heading(
    "Pharma Compliance Analytics Report",
    1
)


document.add_paragraph(
    "Prepared By: Madharapu Chandramouli"
)


document.add_paragraph(
    f"Generated On: "
    f"{datetime.now().strftime('%d-%m-%Y %H:%M')}"
)


# --------------------------
# Dashboard Summary
# --------------------------

document.add_heading(
    "Dashboard Summary",
    level=2
)


document.add_paragraph(
    f"Total Compliance Cases: {total_cases}"
)


document.add_paragraph(
    f"Average Risk Score: {average_risk}"
)


document.add_paragraph(
    f"Critical Cases: {critical_cases}"
)


document.add_paragraph(
    f"Open CAPA: {open_capa}"
)


document.add_paragraph(
    f"Closed CAPA: {closed_capa}"
)


# --------------------------
# Applied Filters
# --------------------------

document.add_heading(
    "Applied Filters",
    level=2
)


document.add_paragraph(
    f"Company: {company}"
)


document.add_paragraph(
    f"Department: {department}"
)


document.add_paragraph(
    f"Plant: {plant}"
)


document.add_paragraph(
    f"Severity: {severity}"
)


# --------------------------
# Management Decision Insights
# --------------------------

document.add_heading(
    "Management Decision Insights",
    level=2
)

document.add_paragraph(
    f"Highest Risk Department: "
    f"{highest_risk_department_display} "
    f"(Risk Score: {highest_risk_score})"
)

document.add_paragraph(
    f"Top Root Cause: {top_root_cause}"
)

document.add_paragraph(
    f"Most Common Deviation: {top_deviation}"
)

document.add_paragraph(
    f"Slowest Closure Department: "
    f"{slowest_department_display} "
    f"(Average Closure: {slowest_closure_days} Days)"
)
# --------------------------
# Recommendations
# --------------------------

document.add_heading(
    "Recommendations",
    level=2
)


document.add_paragraph(
    "• Improve CAPA monitoring and closure processes."
)


document.add_paragraph(
    "• Conduct periodic GMP and SOP training."
)


document.add_paragraph(
    "• Review recurring deviation root causes."
)


document.add_paragraph(
    "• Increase monitoring of high-risk departments "
    "and manufacturing plants."
)


document.add_paragraph(
    "• Use compliance trend analysis to support "
    "preventive management decisions."
)


# ==========================================================
# CREATE WORD FILE
# ==========================================================

buffer = BytesIO()

document.save(buffer)

buffer.seek(0)


st.download_button(
    label="📄 Download Compliance Report",
    data=buffer,
    file_name="Pharma_Compliance_Report.docx",
    mime=(
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.document"
    ),
    width="stretch"
)


# ==========================================================
# DASHBOARD STATUS
# ==========================================================

st.success(
    "Dashboard loaded successfully. "
    "Apply filters to explore compliance insights "
    "and download the report."
)